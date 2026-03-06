import os
import sys
import subprocess

def install_package(package):
    print(f"Instalando {package}...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import docx
except ImportError:
    install_package("python-docx")
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_cover_page(doc, doc_title):
    for _ in range(5):
        doc.add_paragraph()
        
    title = doc.add_heading(doc_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(3):
        doc.add_paragraph()
        
    subject = doc.add_paragraph('Materia:\nGESTIÓN DEL PROCESO DE DESARROLLO DE SOFTWARE')
    subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subject.runs[0].bold = True
    
    for _ in range(3):
        doc.add_paragraph()
        
    team_heading = doc.add_paragraph('Integrantes del Equipo:')
    team_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team_heading.runs[0].bold = True
    
    members = [
        'Cruz Guerrero Dilan Uriel',
        'Ortiz López Fernando',
        'Rodríguez Bueno Jesus Alejandro'
    ]
    
    for member in members:
        p = doc.add_paragraph(member)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_page_break()

def create_estrategias(filepath):
    doc = Document()
    
    add_cover_page(doc, 'Estrategias de Despliegue de Software\n(Proyecto LabManager)')
    
    title = doc.add_heading('Estrategias de Despliegue de Software para LabManager', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Intro
    doc.add_paragraph('Tras revisar los conceptos sobre los métodos para lanzar actualizaciones de software con el objetivo de minimizar fallos e interrupciones, se han seleccionado tres estrategias fundamentales aplicadas al contexto específico del sistema universitario LabManager. La decisión se ha tomado considerando la necesidad de equilibrar el riesgo de afectar a estudiantes y docentes durante sus reservas, optimizando la gestión de inventario en tiempo real.')
    
    doc.add_heading('1. Blue-Green Deployment', level=2)
    doc.add_paragraph('Concepto: Esta estrategia utiliza dos entornos de producción instalados exactamente iguales y en paralelo. Uno es el "Azul", que maneja el tráfico real y estable de nuestro panel de estudiantes, docentes y administradores. El otro es el "Verde", el cual se mantiene en standby y en donde se despliega y prueba la nueva versión de LabManager. Una vez validado, el tráfico de la red universitaria se redirige instantáneamente del Azul al Verde.', style='List Bullet')
    doc.add_paragraph('Justificación de uso para LabManager: Es la estrategia ideal para sistemas críticos que no pueden permitirse tiempos de inactividad, como es el caso de un sistema de reservas de laboratorios donde los estudiantes necesitan equipos para presentar sus proyectos urgentes o clases en vivo. Si una nueva actualización en LabManager llegara a fallar, el rollback es inmediato, permitiendo a los alumnos seguir rentando sus laptops y accesorios desde el entorno Azul intacto.', style='List Bullet')
    
    doc.add_heading('2. Canary Deployment (Lanzamiento Canario)', level=2)
    doc.add_paragraph('Concepto: En esta estrategia, la nueva versión del software no se libera a todos los estudiantes de la universidad al mismo tiempo. En su lugar, se despliega primero de manera exclusiva a un porcentaje muy pequeño de la matrícula (típicamente a un grupo de control de 5% a 10% de alumnos de cierta materia piloto).', style='List Bullet')
    doc.add_paragraph('Justificación de uso para LabManager: Se justifica su implementación porque es sumamente efectivo para validar funciones de alto impacto, como un nuevo algoritmo de validación de códigos QR en el inventario o la nueva bandeja de incidentes. Si la funcionalidad de lectura de códigos QR contiene defectos no detectados, el impacto es minúsculo (solo afecta a 1 o 2 grupos). Esto brinda a los administradores de laboratorios la tranquilidad de que el flujo de préstamos no colapsará masivamente antes de corregirse.', style='List Bullet')

    doc.add_heading('3. Rolling Deployment (Actualización progresiva)', level=2)
    doc.add_paragraph('Concepto: Consiste en actualizar la infraestructura de LabManager de manera progresiva. A medida que ciertos componentes (como el microservicio de notificaciones por WebSockets de Spring Boot) se apagan para actualizarse al nuevo código, el resto de las instancias continúan operando y respondiendo a los alumnos.', style='List Bullet')
    doc.add_paragraph('Justificación de uso para LabManager: La principal ventaja es la rentabilidad de infraestructura. Evitamos a la universidad el alto costo de duplicar sus servidores de aplicaciones y bases de datos al 100%. LabManager mantiene su disponibilidad web y, si el módulo de notificaciones de préstamos sufre de errores puntuales, la caída queda aislada y se revierte solo esa porción de la aplicación sin desconectar a los administradores activamente en el tablero de inventarios.', style='List Bullet')

    doc.save(filepath)
    print(f"Creado: {filepath}")

def create_cicd(filepath):
    doc = Document()
    
    add_cover_page(doc, 'Flujo General de Pipeline CI/CD\n(Proyecto LabManager)')
    
    title = doc.add_heading('Flujo General de Pipeline CI/CD para LabManager', 1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('En este documento se traza el recorrido general que sigue una actualización de nuestro sistema de reservas LabManager, desde la escritura de código, hasta el despliegue a la red universitaria utilizando principios de Integración y Entrega Continua (CI/CD).')
    
    doc.add_heading('Fase 1: Integración y Construcción Continua (CI)', level=2)
    p1 = doc.add_paragraph('', style='List Number')
    p1.add_run('Commit en Repositorio Central: ').bold = True
    p1.add_run('El equipo de desarrollo de LabManager (Front-end en React y Back-end en Spring Boot) envía sus cambios al repositorio utilizando GitLab CI o GitHub Actions.')
    
    p2 = doc.add_paragraph('', style='List Number')
    p2.add_run('Construcción de Imágenes y Artefactos: ').bold = True
    p2.add_run('El pipeline automatizado detecta el push, ejecuta Maven y Vite, corre las pruebas unitarias y empaqueta la aplicación completa en artefactos inmutables (ej. Imágenes de Docker de integradora-frontend e integradora-backend).')
    
    doc.add_heading('Fase 2: Preproducción y Pruebas (Staging)', level=2)
    p3 = doc.add_paragraph('', style='List Number')
    p3.add_run('Despliegue al Laboratorio de Pruebas (Staging): ').bold = True
    p3.add_run('Los contenedores Docker se lanzan a un ambiente clonado. Aquí los testers y docentes de prueba simulan reservas exhaustivas y aperturas de incidentes garantizando que la app rinda al cien.')

    p4 = doc.add_paragraph('', style='List Number')
    p4.add_run('Infraestructura como Código (IaC): ').bold = True
    p4.add_run('Las bases de datos MySQL y la red Docker en staging se levantan automáticamente a través de herramientas automatizadas, asegurando la propiedad de Idempotencia: las pruebas producirán resultados consistentes y aislados sin contaminar la lista real de alumnos.')

    doc.add_heading('Fase 3: Distribución a Producción Universitaria (CD)', level=2)
    p5 = doc.add_paragraph('', style='List Number')
    p5.add_run('Lanzamiento Seguro: ').bold = True
    p5.add_run('Al estar la versión lista y pulida, el código viaja al servidor central del Instituto mediante:\n')
    p5.add_run('• Entrega Continua: ').bold = True
    p5.add_run('El pipeline manda un mensaje, y el desarrollador líder debe presionar "Aprobar Despliegue" de manera manual y controlada.\n')
    p5.add_run('• Despliegue Continuo: ').bold = True
    p5.add_run('Alternativamente, usando herramientas declarativas de despliegue, el update llega instantáneamente al usuario, dictando por ejemplo, la regla de lanzamiento Canario para que solo algunos estudiantes vean el nuevo Dashboard.')

    doc.add_heading('Fase 4: Monitoreo y Tolerancia a Fallas', level=2)
    p6 = doc.add_paragraph('', style='List Number')
    p6.add_run('Evaluación de Métricas DORA: ').bold = True
    p6.add_run('En el ambiente en vivo, los supervisores técnicos revisan las estadísticas de caídas, midiendo la frecuencia y estabilidad de los préstamos en tiempo real con herramientas especializadas.')
    
    p7 = doc.add_paragraph('', style='List Number')
    p7.add_run('Rollback Automático Defensivo: ').bold = True
    p7.add_run('Si una porción vital de código tira el backend (ej. una actualización que rompa la tabla de "laptops"), el sistema de orquestación aborta la misión y recicla la versión exitosa del día anterior para restaurar el servicio inmediatamente a toda la comunidad estudiantil.')

    doc.save(filepath)
    print(f"Creado: {filepath}")

estrategias_path = os.path.abspath("Estrategias_de_Despliegue.docx")
cicd_path = os.path.abspath("Flujo_Pipeline_CICD.docx")

create_estrategias(estrategias_path)
create_cicd(cicd_path)

print("Iniciando conversion a PDF usando Microsoft Word COM...")
ps_script = f'''
$word = New-Object -ComObject Word.Application
$word.Visible = $false

$doc1 = $word.Documents.Open("{estrategias_path}")
$pdf1 = "{estrategias_path.replace(".docx", ".pdf")}"
$doc1.SaveAs([ref] $pdf1, [ref] 17)
$doc1.Close()

$doc2 = $word.Documents.Open("{cicd_path}")
$pdf2 = "{cicd_path.replace(".docx", ".pdf")}"
$doc2.SaveAs([ref] $pdf2, [ref] 17)
$doc2.Close()

$word.Quit()
'''

with open("convert_to_pdf.ps1", "w") as f:
    f.write(ps_script)

result = subprocess.run(["powershell", "-ExecutionPolicy", "Bypass", "-File", "convert_to_pdf.ps1"], capture_output=True, text=True)
if result.returncode == 0:
    print("Documentos convertidos a PDF exitosamente.")
else:
    print("Hubo un error convirtiendo a PDF. El archivo .docx sí fue creado.")
    print(result.stderr)
